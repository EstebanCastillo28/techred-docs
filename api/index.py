from flask import Flask, render_template, request, send_file, jsonify
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import Table, TableStyle
from io import BytesIO
import os, io
from docx import Document
import mammoth
from weasyprint import HTML as WeasyprintHTML

app = Flask(__name__, template_folder="../templates", static_folder="../static")

BASE = os.path.dirname(os.path.abspath(__file__))
PLANTILLAS = os.path.join(BASE, "..", "plantillas")

# ─────────────────────────────────────────────
#  HELPER: rellenar DOCX template → PDF
# ─────────────────────────────────────────────
CSS_BASE = """
@page { size: A4; margin: 18mm 16mm 18mm 16mm; }
body { font-family: Arial, sans-serif; font-size: 9.5pt; color: #000; line-height: 1.5; }
h1 { font-size: 14pt; text-align: center; margin: 3mm 0 2mm; }
h2 { font-size: 11pt; margin: 2mm 0; }
h3 { font-size: 10pt; margin: 2mm 0; }
p  { margin: 1.5mm 0; }
strong { font-weight: bold; }
table { width: 100%; border-collapse: collapse; margin: 3mm 0; }
td, th { border: 0.5pt solid #000; padding: 2mm 3mm; font-size: 9pt; }
th { background: #e0e0e0; font-weight: bold; text-align: center; }
"""

def fill_template_pdf(template_name, data, extra_css=""):
    path = os.path.join(PLANTILLAS, template_name)
    doc = Document(path)

    def repl(text):
        for k, v in data.items():
            text = text.replace(f"{{{{{k}}}}}", str(v or ""))
        return text

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = repl(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = repl(run.text)

    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_buf.seek(0)

    html_body = mammoth.convert_to_html(docx_buf).value
    full_html = f"""<!DOCTYPE html>
<html><head><meta charset='utf-8'>
<style>{CSS_BASE}{extra_css}</style></head>
<body>{html_body}</body></html>"""

    pdf_buf = io.BytesIO()
    WeasyprintHTML(string=full_html).write_pdf(pdf_buf)
    pdf_buf.seek(0)
    return pdf_buf


# ─────────────────────────────────────────────
#  ORDEN DE PEDIDO — ReportLab pixel-perfect
# ─────────────────────────────────────────────
def pdf_orden(d):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4
    m = 15 * mm

    # ── LOGO TECHRED (izquierda)
    c.setFont("Helvetica-Bold", 26)
    c.setFillColor(colors.black)
    c.drawString(m, H - 22*mm, "•TECH")
    c.setFillColor(colors.HexColor("#cc0000"))
    c.drawString(m + 38*mm, H - 22*mm, "RED")
    c.setFillColor(colors.black)

    # Flecha del logo
    c.setFont("Helvetica-Bold", 14)
    c.drawString(m + 63*mm, H - 22*mm, "⇌")

    c.setFont("Helvetica", 8.5)
    c.drawString(m, H - 28*mm, "Tecnología  •  Infraestructura")
    c.setFont("Helvetica-Bold", 8)
    c.drawString(m, H - 34*mm, "NIT. 900.866.525 -6")
    c.drawString(m, H - 39*mm, "CLL 68 # 23-61 7 DE AGOSTO")
    c.drawString(m, H - 44*mm, "CORREO: VANTILISTO@TECHRED.COM.CO")
    c.drawString(m, H - 49*mm, "www.techred.com.co")

    # ── CAMPOS CLIENTE (centro)
    cx = 78*mm
    def lf(label, val, x, y, lw=22*mm, fw=80*mm, sz=9):
        c.setFont("Helvetica-Bold", sz)
        c.drawString(x, y, label)
        c.setFont("Helvetica", sz)
        c.drawString(x + lw, y, str(val or ""))
        c.setLineWidth(0.4)
        c.line(x + lw - 1, y - 2.5, x + lw + fw, y - 2.5)

    lf("CLIENTE:", d.get("cliente"), cx, H-20*mm, 20*mm, 92*mm)
    lf("C.C:",    d.get("cc"),      cx, H-28*mm, 11*mm, 34*mm)
    lf("TEL/CEL:", d.get("tel"),    cx+48*mm, H-28*mm, 19*mm, 42*mm)
    lf("DIRECCIÓN:", d.get("direccion"), cx, H-36*mm, 22*mm, 90*mm)
    lf("BARRIO:",  d.get("barrio"), cx, H-44*mm, 16*mm, 46*mm)
    lf("CIUDAD:",  d.get("ciudad"), cx+66*mm, H-44*mm, 16*mm, 27*mm)
    lf("CONTACTO 1:", d.get("contacto1"), cx, H-52*mm, 26*mm, 86*mm)
    lf("CONTACTO 2:", d.get("contacto2"), cx, H-60*mm, 26*mm, 86*mm)
    lf("CORREO:",  d.get("correo"), cx, H-68*mm, 16*mm, 96*mm)

    # ── CAJA ORDEN / FECHA / CUENTA (derecha)
    bx = 168*mm
    # Caja ORDEN DE PEDIDO
    c.setLineWidth(1.2)
    c.roundRect(bx, H-29*mm, 27*mm, 10*mm, 2)
    c.setFont("Helvetica-Bold", 9)
    c.setFillColor(colors.HexColor("#cc0000"))
    c.drawCentredString(bx+13.5*mm, H-23.5*mm, "ORDEN DE")
    c.drawCentredString(bx+13.5*mm, H-27.5*mm, "PEDIDO")
    c.setFillColor(colors.black)

    # FECHA DE PEDIDO
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(bx, H-37*mm, "FECHA DE PEDIDO")
    date_labels = ["DIA","MES","AÑO"]
    date_vals   = [d.get("dia",""), d.get("mes",""), d.get("anio","")]
    date_widths = [8*mm, 8*mm, 11*mm]
    dx = bx
    for lbl, val, dw in zip(date_labels, date_vals, date_widths):
        c.setLineWidth(0.6)
        c.roundRect(dx, H-47*mm, dw, 8*mm, 1)
        c.setFont("Helvetica-Bold", 5.5)
        c.drawCentredString(dx + dw/2, H-41.5*mm, lbl)
        c.setFont("Helvetica", 8)
        c.drawCentredString(dx + dw/2, H-45.5*mm, str(val))
        dx += dw + 0.5*mm

    # CUENTA CONTRATO
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(bx, H-55*mm, "CUENTA CONTRATO")
    c.setLineWidth(0.6)
    c.roundRect(bx, H-65*mm, 27*mm, 9*mm, 2)
    c.setFont("Helvetica", 9)
    c.drawCentredString(bx+13.5*mm, H-60.5*mm, str(d.get("cuenta_contrato","")))

    # ── TABLA PRODUCTOS
    items = d.get("items", [])
    while len(items) < 4: items.append({})
    rows = [["REF","CANT","DESCRIPCION","VALOR PRODUCTO","CUOTAS","VALOR CUOTAS"]]
    for it in items[:4]:
        rows.append([
            it.get("ref",""), it.get("cant",""), it.get("descripcion",""),
            it.get("valor_producto",""), it.get("cuotas",""), it.get("valor_cuotas","")
        ])
    rows.append(["","","VALOR TOTAL DE LA COMPRA", f"$ {d.get('valor_total','')}","",""])

    col_w = [22*mm, 13*mm, 79*mm, 30*mm, 17*mm, 33*mm]
    tbl = Table(rows, colWidths=col_w, rowHeights=[8*mm]+[11*mm]*4+[9*mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0), colors.black),
        ("TEXTCOLOR",(0,0),(-1,0),  colors.white),
        ("FONTNAME",(0,0),(-1,0),   "Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),  8),
        ("GRID",(0,0),(-1,-1),      0.5, colors.black),
        ("ALIGN",(0,0),(-1,-1),     "CENTER"),
        ("VALIGN",(0,0),(-1,-1),    "MIDDLE"),
        ("ALIGN",(2,1),(2,-1),      "LEFT"),
        ("SPAN",(2,5),(5,5)),
        ("BACKGROUND",(0,5),(-1,5), colors.black),
        ("TEXTCOLOR",(0,5),(-1,5),  colors.white),
        ("FONTNAME",(2,5),(5,5),    "Helvetica-Bold"),
    ]))
    tbl_y = H - 130*mm
    tbl.wrapOn(c, W, H)
    tbl.drawOn(c, m, tbl_y)

    # ── FORMA DE PAGO
    py = tbl_y - 16*mm
    c.setFont("Helvetica-Bold", 10)
    c.drawString(m, py, "FORMA DE PAGO")
    pays = [("EFECTIVO:","efectivo"),("VANTI LISTO:","vanti_listo"),("OTRO","otro")]
    px = m
    for lbl, key in pays:
        c.setFont("Helvetica-Bold", 9)
        c.drawString(px, py-10*mm, lbl)
        c.setLineWidth(0.8)
        c.rect(px+26*mm, py-13*mm, 7*mm, 7*mm)
        if d.get(key):
            c.setFont("Helvetica-Bold", 11)
            c.drawCentredString(px+29.5*mm, py-8.2*mm, "X")
        px += 42*mm

    lf("CONTRAENTREGA: $", d.get("contraentrega"), m, py-24*mm, 38*mm, 50*mm)
    lf("ASESOR:",          d.get("asesor"),         m, py-33*mm, 18*mm, 70*mm)
    lf("COD ASESOR:",      d.get("cod_asesor"),     m, py-42*mm, 26*mm, 62*mm)
    lf("SUPERVISOR:",      d.get("supervisor"),     m, py-51*mm, 24*mm, 64*mm)

    # Caja huella
    hx, hy = 130*mm, py-58*mm
    c.setLineWidth(0.8)
    c.roundRect(hx, hy, 26*mm, 38*mm, 4)
    c.setFont("Helvetica-Bold", 8)
    c.drawCentredString(hx+13*mm, hy - 5*mm, "HUELLA CLIENTE")

    # Firma cliente
    c.line(170*mm, hy, 205*mm, hy)
    c.setFont("Helvetica-Bold", 8)
    c.drawCentredString(187.5*mm, hy-5*mm, "FIRMA CLIENTE")

    # ── FOOTER
    c.setFillColor(colors.HexColor("#cc0000"))
    c.setFont("Helvetica-Bold", 8.5)
    c.drawCentredString(W/2, 24*mm,
        "Línea de servicio al cliente y garantías - WhatsApp: - 302 5555083 - correo electrónico: servicio.cliente@techred.com.co")
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 7.5)
    c.drawCentredString(W/2, 18*mm, "Conozca más en www.techred.com.co/pages/politicas-de-privacidad")
    c.setFont("Helvetica", 7)
    c.drawCentredString(W/2, 12*mm,
        "Aplican T&C. Este producto es comercializado y distribuido por Techred S.A.S Vanti Listo actúa únicamente como medio de financiamiento. © Techred S.A.S 2026")

    c.showPage()
    c.save()
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
#  ROUTES
# ─────────────────────────────────────────────
@app.route("/")
def home():
    return render_template("index.html")


@app.route("/generate/orden-pedido", methods=["POST"])
def gen_orden():
    data = request.get_json() or {}
    return send_file(pdf_orden(data), as_attachment=True,
                     download_name="orden-pedido.pdf", mimetype="application/pdf")


@app.route("/generate/acta-entrega", methods=["POST"])
def gen_acta():
    data = request.get_json() or {}
    items = data.get("items", [])
    flat = {**data}
    for i, item in enumerate(items[:7], 1):
        flat[f"DESC_{i}"]  = item.get("descripcion","")
        flat[f"CANT_{i}"]  = item.get("cantidad","")
        flat[f"VALOR_{i}"] = item.get("valor","")
    for i in range(len(items)+1, 8):
        flat[f"DESC_{i}"] = flat[f"CANT_{i}"] = flat[f"VALOR_{i}"] = ""
    pdf = fill_template_pdf("acta_entrega_template.docx", flat)
    return send_file(pdf, as_attachment=True,
                     download_name="acta-entrega.pdf", mimetype="application/pdf")


@app.route("/generate/cambio-producto", methods=["POST"])
def gen_cambio():
    data = request.get_json() or {}
    pdf = fill_template_pdf("cambio_producto_template.docx", data)
    return send_file(pdf, as_attachment=True,
                     download_name="cambio-producto.pdf", mimetype="application/pdf")


if __name__ == "__main__":
    app.run(debug=True)
